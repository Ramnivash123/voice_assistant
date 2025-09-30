import re
import os
import tempfile
from datetime import datetime
import base64
import io

import docx
from docx import Document

import numpy as np
import whisper
from gtts import gTTS
import pygame
import threading
import time
from datetime import datetime, timedelta

import streamlit as st
from streamlit_audio_recorder import st_audio_recorder

# -----------------------------
# Timer Thread
# -----------------------------
class ExamTimer(threading.Thread):
    def __init__(self, duration_seconds=7200):
        super().__init__(daemon=True)
        self.start_time = datetime.now()
        self.end_time = self.start_time + timedelta(seconds=duration_seconds)
        self._is_running = True

    def remaining_time(self):
        if not self._is_running:
            return 0
        now = datetime.now()
        if now >= self.end_time:
            return 0
        return int((self.end_time - now).total_seconds())

    def formatted_remaining(self):
        seconds = self.remaining_time()
        if seconds <= 0:
            return "Time is up."
        mins, secs = divmod(seconds, 60)
        hrs, mins = divmod(mins, 60)
        return f"{hrs:02d}:{mins:02d}:{secs:02d} remaining"

    def stop(self):
        self._is_running = False

# -----------------------------
# Config
# -----------------------------
INPUT_DOC = "mugilanQp.docx"
OUTPUT_DOC = "answers.docx"
RECORD_SECONDS = 10
SAMPLE_RATE = 16000
WHISPER_MODEL = "base"

EXCLUDE_SUBSTRS = {
    "answer all questions",
    "name & signature",
    "department of data science",
    "max. marks",
    "time duration",
    "affiliated to",
    "college",
    "batch:", "class:", "subject title:", "semester:",
    "mid term", "reviewer"
}

# -----------------------------
# Audio Helpers
# -----------------------------
def speak_text(text: str):
    """Convert text to speech and play with pygame."""
    if not pygame.mixer.get_init():
        pygame.mixer.init()
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as fp:
        temp_mp3 = fp.name
    try:
        gTTS(text=text, lang="en").save(temp_mp3)
        pygame.mixer.music.load(temp_mp3)
        pygame.mixer.music.play()
        clock = pygame.time.Clock()
        while pygame.mixer.music.get_busy():
            clock.tick(10)
    finally:
        try:
            pygame.mixer.music.unload()
        except Exception:
            pass
        if os.path.exists(temp_mp3):
            os.remove(temp_mp3)

def save_audio_data(audio_data, path: str):
    """Save audio data from streamlit-audio-recorder to WAV file."""
    try:
        # If audio_data is bytes, write directly to file
        if isinstance(audio_data, bytes):
            with open(path, "wb") as f:
                f.write(audio_data)
        else:
            st.error("Invalid audio data format")
    except Exception as e:
        st.error(f"Error saving audio: {e}")

def transcribe_audio(audio_data, model) -> str:
    """Transcribe audio data with Whisper."""
    try:
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_file:
            temp_path = temp_file.name
        
        # Save audio data
        save_audio_data(audio_data, temp_path)
        
        # Transcribe
        result = model.transcribe(temp_path, fp16=False)
        
        # Clean up
        if os.path.exists(temp_path):
            os.remove(temp_path)
            
        return (result.get("text") or "").strip()
        
    except Exception as e:
        st.error(f"Transcription error: {e}")
        return ""

# -----------------------------
# Docx Extractor
# -----------------------------
def get_all_text(doc_path):
    """Extract text from both paragraphs and tables."""
    doc = docx.Document(doc_path)
    lines = []

    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    lines.append(txt)

    return lines

def clean_line(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip())

def is_excluded(line: str) -> bool:
    l = line.lower()
    return any(key in l for key in EXCLUDE_SUBSTRS)

def extract_questions(path: str):
    """Extract questions from both paragraphs + tables."""
    raw_lines = get_all_text(path)
    lines, prev = [], None
    for l in raw_lines:
        l = clean_line(l)
        if l and l != prev:
            lines.append(l)
        prev = l

    section, result, i = None, [], 0

    while i < len(lines):
        line = lines[i]

        if line.lower().startswith("section a"):
            section = "A"; i += 1; continue
        if line.lower().startswith("section b"):
            section = "B"; i += 1; continue
        if line.lower().startswith("section c"):
            section = "C"; i += 1; continue

        if not section or is_excluded(line):
            i += 1
            continue

        if section == "A":
            m = re.match(r"^(\d{1,2})\s+(.*)$", line)
            if m:
                qnum, stem_text = m.groups()
                i += 1
            elif re.fullmatch(r"\d{1,2}", line):
                qnum = line
                i += 1
                stem_text = ""
                if i < len(lines) and not re.match(r"^[ABCD]\b", lines[i], re.I):
                    stem_text = lines[i].strip()
                    i += 1
            else:
                i += 1
                continue

            options = {}
            while i < len(lines) and re.match(r"^[ABCD]\b", lines[i], re.I):
                letter = lines[i][0].upper()
                value = lines[i][1:].strip()
                options[letter] = value
                i += 1

            text = stem_text
            for letter in ["A", "B", "C", "D"]:
                if letter in options:
                    text += f"\n{letter}. {options[letter]}"

            result.append({"section": "A", "label": qnum, "text": text})
            continue

        m = re.match(r"^(\d{1,2})\s*([AB])?\s*(.*)$", line)
        if section in {"B","C"} and m:
            qnum, ab, rest = m.groups(); ab = ab or ""
            block = [rest] if rest else []; i += 1
            while i < len(lines) and not re.match(r"^\d{1,2}\s*[AB]?", lines[i]) and not lines[i].lower().startswith("section"):
                if not is_excluded(lines[i]) and lines[i] != "(OR)":
                    if lines[i] in {"A","B"} and not block:
                        ab = lines[i]
                    else:
                        block.append(lines[i])
                i += 1
            text = " ".join(block)
            result.append({"section": section, "label": f"{qnum} {ab}".strip(), "text": text})
            continue

        i += 1

    return result

# -----------------------------
# Save Answers to DOCX
# -----------------------------
def save_answers_docx(path: str, qa_items):
    doc = Document()
    doc.add_heading("Answers Document", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    for item in qa_items:
        doc.add_paragraph(f"Q{item['label']}: {item['text']}", style="List Bullet")
        doc.add_paragraph(f"A{item['label']}: {item.get('answer','').strip()}\n")
    doc.save(path)

# -----------------------------
# Extract Metadata
# -----------------------------
def extract_metadata(path: str):
    """Extract name and subject title from the DOCX."""
    raw_lines = get_all_text(path)
    name, subject = None, None
    for line in raw_lines:
        if line.strip().lower().startswith("name:"):
            name = line.split(":", 1)[-1].strip()
        if line.strip().lower().startswith("subject title:"):
            subject = line.split(":", 1)[-1].strip()
    return name, subject

# -----------------------------
# Streamlit UI
# -----------------------------
def main():
    st.set_page_config(page_title="Voice Assistant", layout="wide")
    
    st.title("Voice Assistant")
    st.markdown("---")
    
    # Initialize session state
    if 'timer' not in st.session_state:
        st.session_state.timer = None
    if 'exam_started' not in st.session_state:
        st.session_state.exam_started = False
    if 'current_question' not in st.session_state:
        st.session_state.current_question = 0
    if 'qa_items' not in st.session_state:
        st.session_state.qa_items = []
    if 'questions' not in st.session_state:
        st.session_state.questions = []
    if 'model' not in st.session_state:
        st.session_state.model = None
    if 'audio_data' not in st.session_state:
        st.session_state.audio_data = None

    # Sidebar for controls
    with st.sidebar:
        st.header("Exam Controls")
        
        if st.session_state.timer:
            remaining = st.session_state.timer.formatted_remaining()
            st.metric("Time Remaining", remaining)
        
        if st.button("🔄 Refresh Timer"):
            if st.session_state.timer:
                st.rerun()
        
        st.markdown("---")
        st.header("Questions Progress")
        if st.session_state.questions:
            total = len(st.session_state.questions)
            current = st.session_state.current_question
            st.progress(current / total if total > 0 else 0)
            st.write(f"Question {current}/{total}")

    # Main area
    if not st.session_state.exam_started:
        show_welcome_screen()
    else:
        show_exam_interface()

def show_welcome_screen():
    st.header("Welcome to Voice Assistant Exam")
    
    uploaded_file = st.file_uploader("Upload Question Paper (DOCX)", type=["docx"])
    
    if uploaded_file is not None:
        # Save uploaded file
        with open(INPUT_DOC, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        # Extract metadata
        name, subject = extract_metadata(INPUT_DOC)
        
        if name and subject:
            st.success(f"✅ Question paper loaded successfully!")
            st.write(f"**Candidate:** {name}")
            st.write(f"**Subject:** {subject}")
            
            if st.button("Start Exam", type="primary"):
                start_exam(name, subject)
        else:
            st.error("Could not extract name and subject from the question paper.")

def start_exam(name, subject):
    """Initialize exam components"""
    try:
        # Load Whisper model
        with st.spinner("Loading speech recognition model..."):
            st.session_state.model = whisper.load_model(WHISPER_MODEL)
        
        # Extract questions
        with st.spinner("Extracting questions from document..."):
            st.session_state.questions = extract_questions(INPUT_DOC)
        
        if not st.session_state.questions:
            st.error("No questions found in the document!")
            return
        
        # Initialize timer
        st.session_state.timer = ExamTimer(duration_seconds=7200)
        st.session_state.timer.start()
        
        st.session_state.exam_started = True
        st.session_state.current_question = 0
        st.session_state.qa_items = []
        
        st.rerun()
        
    except Exception as e:
        st.error(f"Error starting exam: {str(e)}")

def show_exam_interface():
    """Display the main exam interface"""
    # Timer display at top
    if st.session_state.timer:
        remaining = st.session_state.timer.formatted_remaining()
        if "Time is up" in remaining:
            st.error("⏰ Time's up! Exam completed.")
            complete_exam()
            return
    
    # Show current question
    if st.session_state.current_question < len(st.session_state.questions):
        show_current_question()
    else:
        complete_exam()

def show_current_question():
    """Display and handle the current question"""
    idx = st.session_state.current_question
    q = st.session_state.questions[idx]
    
    st.header(f"Question {idx + 1}")
    
    # Automatically read question aloud when first displayed
    if f'question_{idx}_read' not in st.session_state:
        try:
            speak_text(f"Question {idx + 1}. {q['text']}")
            st.session_state[f'question_{idx}_read'] = True
        except Exception as e:
            st.error(f"Audio playback error: {e}")
    
    # Question display
    with st.expander("View Question", expanded=True):
        st.write(q['text'])
    
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col1:
        if st.button("🔊 Repeat Question", use_container_width=True):
            try:
                speak_text(f"Question {idx + 1}. {q['text']}")
            except Exception as e:
                st.error(f"Audio playback error: {e}")
    
    with col2:
        if st.button("⏭️ Skip Question", use_container_width=True):
            st.session_state.qa_items.append({
                "label": str(idx + 1), 
                "text": q['text'], 
                "answer": "[SKIPPED]"
            })
            st.session_state.current_question += 1
            st.rerun()
    
    with col3:
        if st.button("⏰ Check Time", use_container_width=True):
            if st.session_state.timer:
                remaining = st.session_state.timer.formatted_remaining()
                st.info(f"Time remaining: {remaining}")
                try:
                    speak_text(remaining)
                except Exception as e:
                    st.error(f"Audio playback error: {e}")
    
    st.markdown("---")
    st.subheader("Record Your Answer")
    
    # Audio recorder component
    audio_data = st_audio_recorder()
    
    if audio_data is not None:
        st.session_state.audio_data = audio_data
        if st.button("🎙️ Process Recording", type="primary"):
            process_audio_answer(idx, q)

def process_audio_answer(idx, q):
    """Process the recorded audio answer"""
    if st.session_state.audio_data is None:
        st.error("No audio recorded. Please record your answer first.")
        return
    
    try:
        with st.spinner("Transcribing your answer..."):
            answer = transcribe_audio(st.session_state.audio_data, st.session_state.model).lower()
        
        # Handle special commands
        if "skip" in answer:
            st.info("Question skipped based on your voice command")
            st.session_state.qa_items.append({
                "label": str(idx + 1), 
                "text": q['text'], 
                "answer": "[SKIPPED]"
            })
            st.session_state.current_question += 1
            st.session_state.audio_data = None
            st.rerun()
            
        elif "repeat" in answer:
            st.info("Repeating question based on your voice command...")
            try:
                speak_text(f"Question {idx + 1}. {q['text']}")
            except Exception as e:
                st.error(f"Audio playback error: {e}")
            # Clear audio data for new recording
            st.session_state.audio_data = None
            st.rerun()
            
        else:
            # Normal answer
            st.session_state.qa_items.append({
                "label": str(idx + 1), 
                "text": q['text'], 
                "answer": answer
            })
            
            st.success("Answer recorded successfully!")
            st.write(f"**Your answer:** {answer}")
            
            # Read back the answer
            try:
                speak_text(f"You answered: {answer}")
            except Exception as e:
                st.error(f"Audio playback error: {e}")
            
            # Clear audio data and move to next question
            st.session_state.audio_data = None
            st.session_state.current_question += 1
            st.rerun()
        
    except Exception as e:
        st.error(f"Error processing answer: {str(e)}")

def complete_exam():
    """Complete the exam and show results"""
    st.header("🎉 Exam Completed!")
    
    # Save answers document
    if st.session_state.qa_items:
        save_answers_docx(OUTPUT_DOC, st.session_state.qa_items)
        st.success(f"Answers saved to: {OUTPUT_DOC}")
        
        # Provide download link
        with open(OUTPUT_DOC, "rb") as file:
            st.download_button(
                label="📥 Download Answers Document",
                data=file,
                file_name=OUTPUT_DOC,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    # Show summary
    st.subheader("Exam Summary")
    total_questions = len(st.session_state.questions)
    answered = len([q for q in st.session_state.qa_items if q.get('answer') not in ['[SKIPPED]', '']])
    skipped = total_questions - answered
    
    st.write(f"**Total Questions:** {total_questions}")
    st.write(f"**Questions Answered:** {answered}")
    st.write(f"**Questions Skipped:** {skipped}")
    
    if st.button("Start New Exam"):
        # Reset session state
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

if __name__ == "__main__":
    main()
